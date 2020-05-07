#!/usr/bin/python3
# -*- coding: utf-8 -*-
# @Time    :2019/04/30 01:33
# @Author  : Qdan
# @File    :word.py
# @Project :PyCharm
# @Description: ::

from pdfminer.pdfparser import PDFParser, PDFDocument
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.layout import LAParams
from pdfminer.converter import PDFPageAggregator
from pdfminer.pdfinterp import PDFTextExtractionNotAllowed

from docx import Document

import warnings
warnings.filterwarnings("ignore")

import os

def pdf_to_docx(source_file, docx_src, docx_name):
    """
    将指定的 PDF 文格式件转换为 docx 格式
    :param source_file: 源文件
    :param docx_src: docx 新文件存储路径
    :param docx_name: 新文件命名
    :return: null 但有文件处理完成提示
    """
    document = Document()
    fn = open(source_file, 'rb')
    parser = PDFParser(fn)
    doc = PDFDocument()

    parser.set_document(doc)
    doc.set_parser(parser)
    resource = PDFResourceManager()
    laparams = LAParams()
    device = PDFPageAggregator(resource, laparams=laparams)
    interpreter = PDFPageInterpreter(resource, device)

    for i in doc.get_pages():
        interpreter.process_page(i)
        layout = device.get_result()
        for out in layout:
            if hasattr(out, "get_text"):
                content = out.get_text().replace(u'\xa0', u' ')
                document.add_paragraph(
                    content, style='ListBullet'
                )
            docx_src = docx_src
            if not os.path.exists(docx_src):
                os.makefiles(docx_src)
            docx_name = docx_name
            document.save(f'{docx_src}{docx_name}' + '.docx')
    print(f'-- {source_file} --处理为 docx 完成')

def traversing_folder(text):
    """
    遍历当前文件夹里面的所有文件
    :param text: 遍历文件夹路径
    :return: null
    """
    files = os.listdir(text)
    for file in files:
        source_file = text + '\\' + file
        print(file)
        # print(source_file)
        file = file.replace(".pdf", "")
        pdf_to_docx(source_file, r"./datas/docx/", file)

traversing_folder(r"./datas/pdf/")
